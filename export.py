"""
WriterFlow — Export Service
============================
PDF, DOCX, and EPUB generation.
Imports BookService / ChapterService from database.py.
"""
import io
from utils import strip_markdown


# ─── DOCX ────────────────────────────────────────────────────────────────────

def export_to_docx(book_id: int) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from database import BookService, ChapterService

    book     = BookService().get_book(book_id)
    chapters = ChapterService().get_chapters(book_id)

    doc = Document()

    # ── Cover image ──────────────────────────────────────────────────────────
    if book.get("cover_image"):
        cover_buf = io.BytesIO(book["cover_image"])
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(cover_buf, width=Inches(3.5))
        except Exception:
            pass
        doc.add_paragraph()

    # ── Title page ───────────────────────────────────────────────────────────
    title_para = doc.add_heading(book["title"], 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if book.get("genre"):
        p = doc.add_paragraph(f"Gênero: {book['genre']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if book.get("synopsis"):
        doc.add_paragraph()
        doc.add_heading("Sinopse", level=2)
        doc.add_paragraph(book["synopsis"])

    doc.add_page_break()

    # ── Table of contents ────────────────────────────────────────────────────
    doc.add_heading("Sumário", level=1)
    for i, ch in enumerate(chapters, 1):
        p = doc.add_paragraph(f"{i}. {ch['title']}", style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
    doc.add_page_break()

    # ── Chapters ─────────────────────────────────────────────────────────────
    for ch in chapters:
        doc.add_heading(ch["title"], level=1)
        content = strip_markdown(ch.get("content") or "")
        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── PDF ─────────────────────────────────────────────────────────────────────

def export_to_pdf(book_id: int) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        HRFlowable, Image,
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from database import BookService, ChapterService

    book     = BookService().get_book(book_id)
    chapters = ChapterService().get_chapters(book_id)

    buf = io.BytesIO()
    W, H = A4
    margin = 3 * cm

    def _add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.HexColor("#999999"))
        canvas.drawCentredString(W / 2, 1.5 * cm, str(canvas.getPageNumber()))
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=margin, leftMargin=margin,
        topMargin=margin, bottomMargin=2.5 * cm,
        onFirstPage=_add_page_number,
        onLaterPages=_add_page_number,
    )

    styles = getSampleStyleSheet()
    S = {
        "title": ParagraphStyle(
            "BTitle", parent=styles["Title"],
            fontSize=30, spaceAfter=18, leading=36,
            textColor=colors.HexColor("#1a1a2e"), alignment=TA_CENTER,
        ),
        "subtitle": ParagraphStyle(
            "BSub", parent=styles["Normal"],
            fontSize=13, spaceAfter=8, textColor=colors.HexColor("#555"),
            alignment=TA_CENTER,
        ),
        "synopsis_label": ParagraphStyle(
            "BSlbl", parent=styles["Normal"],
            fontSize=11, fontName="Helvetica-Bold",
            spaceBefore=16, spaceAfter=6, textColor=colors.HexColor("#333"),
        ),
        "synopsis": ParagraphStyle(
            "BSyn", parent=styles["Normal"],
            fontSize=11, leading=17, spaceAfter=8,
            textColor=colors.HexColor("#444"),
        ),
        "chapter": ParagraphStyle(
            "BCh", parent=styles["Heading1"],
            fontSize=22, spaceAfter=14, spaceBefore=8,
            textColor=colors.HexColor("#1a1a2e"), alignment=TA_LEFT,
        ),
        "body": ParagraphStyle(
            "BBody", parent=styles["Normal"],
            fontSize=12, leading=19, spaceAfter=9,
            textColor=colors.HexColor("#222"), alignment=TA_JUSTIFY,
            firstLineIndent=18,
        ),
        "toc_title": ParagraphStyle(
            "BToc", parent=styles["Heading1"],
            fontSize=18, spaceAfter=12,
        ),
        "toc_entry": ParagraphStyle(
            "BTocE", parent=styles["Normal"],
            fontSize=11, leading=16, spaceAfter=4,
            leftIndent=12, textColor=colors.HexColor("#333"),
        ),
    }

    story = []

    # ── Cover image ──────────────────────────────────────────────────────────
    if book.get("cover_image"):
        try:
            cover_buf = io.BytesIO(book["cover_image"])
            img = Image(cover_buf, width=10 * cm, height=15 * cm, kind="proportional")
            img.hAlign = "CENTER"
            story.append(Spacer(1, 1 * cm))
            story.append(img)
            story.append(Spacer(1, 1 * cm))
        except Exception:
            pass

    # ── Title page ───────────────────────────────────────────────────────────
    story.append(Paragraph(book["title"], S["title"]))
    if book.get("genre"):
        story.append(Paragraph(f"Gênero: {book['genre']}", S["subtitle"]))
    if book.get("synopsis"):
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph("Sinopse", S["synopsis_label"]))
        story.append(Paragraph(book["synopsis"], S["synopsis"]))
    story.append(PageBreak())

    # ── Table of contents ────────────────────────────────────────────────────
    story.append(Paragraph("Sumário", S["toc_title"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#ccc")))
    story.append(Spacer(1, 0.3 * cm))
    for i, ch in enumerate(chapters, 1):
        story.append(Paragraph(f"{i}.  {ch['title']}", S["toc_entry"]))
    story.append(PageBreak())

    # ── Chapters ─────────────────────────────────────────────────────────────
    for ch in chapters:
        story.append(Paragraph(ch["title"], S["chapter"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#ddd")))
        story.append(Spacer(1, 0.3 * cm))
        content = strip_markdown(ch.get("content") or "")
        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para.replace("\n", "<br/>"), S["body"]))
        story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


# ─── EPUB ─────────────────────────────────────────────────────────────────────

def export_to_epub(book_id: int) -> bytes:
    from ebooklib import epub
    import markdown2
    from database import BookService, ChapterService

    book     = BookService().get_book(book_id)
    chapters = ChapterService().get_chapters(book_id)

    ebook = epub.EpubBook()
    ebook.set_title(book["title"])
    ebook.set_language("pt")
    ebook.add_metadata("DC", "description", book.get("synopsis") or "")
    ebook.add_metadata("DC", "subject",     book.get("genre") or "")

    # ── Cover image ──────────────────────────────────────────────────────────
    if book.get("cover_image"):
        try:
            ebook.set_cover("cover.jpg", book["cover_image"])
        except Exception:
            pass

    # ── Stylesheet ───────────────────────────────────────────────────────────
    css_content = """
        body  { font-family: Georgia, serif; font-size: 1em;
                line-height: 1.7; margin: 1.5em 2em; color: #222; }
        h1    { font-size: 1.9em; margin: 0 0 1em; text-align: center;
                border-bottom: 1px solid #ddd; padding-bottom: 0.5em; }
        h2,h3 { font-size: 1.3em; margin: 1.5em 0 0.6em; }
        p     { margin: 0 0 0.8em; text-align: justify; text-indent: 1.5em; }
        p:first-of-type { text-indent: 0; }
        blockquote { border-left: 3px solid #aaa; padding-left: 1em;
                     font-style: italic; color: #555; margin: 1em 0; }
        pre,code { font-family: monospace; background: #f4f4f4;
                   padding: 0.2em 0.4em; border-radius: 3px; }
    """
    css = epub.EpubItem(
        uid="style", file_name="style.css",
        media_type="text/css", content=css_content,
    )
    ebook.add_item(css)

    # ── Title / synopsis page ────────────────────────────────────────────────
    synopsis_html = ""
    if book.get("synopsis"):
        synopsis_html = f"""
        <div style="margin-top:2em;padding:1em;background:#f9f9f9;border-radius:4px">
            <h2>Sinopse</h2>
            <p>{book['synopsis']}</p>
        </div>"""
    cover_page = epub.EpubHtml(title=book["title"], file_name="title.xhtml", lang="pt")
    cover_page.content = f"""
        <html><body>
        <h1 style="text-align:center;margin-top:3em">{book["title"]}</h1>
        {"<p style='text-align:center;color:#666'>" + book['genre'] + "</p>" if book.get('genre') else ""}
        {synopsis_html}
        </body></html>"""
    cover_page.add_item(css)
    ebook.add_item(cover_page)

    # ── Chapters ─────────────────────────────────────────────────────────────
    epub_chapters = [cover_page]
    toc = []

    for i, ch in enumerate(chapters, 1):
        html_body = markdown2.markdown(
            ch.get("content") or "",
            extras=["fenced-code-blocks", "tables", "strike", "footnotes"],
        )
        ec = epub.EpubHtml(
            title=ch["title"],
            file_name=f"chapter_{i:03d}.xhtml",
            lang="pt",
        )
        ec.content = f"<html><body><h1>{ch['title']}</h1>{html_body}</body></html>"
        ec.add_item(css)
        ebook.add_item(ec)
        epub_chapters.append(ec)
        toc.append(epub.Link(f"chapter_{i:03d}.xhtml", ch["title"], f"ch{i}"))

    ebook.toc   = toc
    ebook.add_item(epub.EpubNcx())
    ebook.add_item(epub.EpubNav())
    ebook.spine = ["nav"] + epub_chapters

    buf = io.BytesIO()
    epub.write_epub(buf, ebook)
    return buf.getvalue()
